from docx import Document
from datetime import datetime

# Function to replace multiple placeholders in both paragraphs and table cells
def replace_placeholders(doc_path, replacements):
    # Open the document
    doc = Document(doc_path)
    
    # Replace placeholders in regular paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            if placeholder in paragraph.text:
                # Replace the placeholder with the corresponding replacement
                paragraph.text = paragraph.text.replace(placeholder, replacement)
    
    # Replace placeholders in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Replace placeholders in each paragraph of the cell
                for paragraph in cell.paragraphs:
                    for placeholder, replacement in replacements.items():
                        if placeholder in paragraph.text:
                            # Replace the placeholder with the corresponding replacement
                            paragraph.text = paragraph.text.replace(placeholder, replacement)
    
    # Extract values for job_position and company_name from replacements
    job_position = replacements.get('{job_position}', 'job_position')
    company_name = replacements.get('{company_name}', 'company_name')
    fileName_date = datetime.today().strftime('%y%m%d')
    
    # Save the modified document with the new file name
    new_file_name = f"{fileName_date + " resume_"+job_position}_{company_name}.docx"  # Use formatted string for naming
    doc.save(new_file_name)
    print(f"Document saved as '{new_file_name}'.")

# Example usage
doc_path = "Resume_template.docx"  # Path to your .docx file
today = datetime.today().strftime('%B %d, %Y')

# Define the placeholders and their corresponding replacements
replacements = {
    '{job_position}': input("Enter the replacement for {job_position}: "),
    '{company_name}': input("Enter the replacement for {company_name}: "),
    '{company_location}': input("Enter the replacement for {company_location}: "),
    '{company_mission}': input("Enter the replacement for {company_mission}: "),
    '{date}': today
}

replace_placeholders(doc_path, replacements)
